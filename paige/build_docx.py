#!/usr/bin/env python3
"""Build a KDP-ready DOCX for PAIGE: Volume One from the chapter markdown files.

Spec (per production brief):
  Trim 5.5 x 8.5in | Garamond body | first-line indent 0.30in | justified body
  Inside margin 0.75in | outside 0.50in | top/bottom 0.75in (mirror margins)
  Widow/orphan control | chapters start on a new page | restarting page numbers
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
CH_DIR = ROOT / "chapters"
OUT = ROOT / "Paige-Volume-One.docx"

AUTHOR = "Kerron Pierre"      # editable
YEAR = "2026"
BODY_FONT = "Garamond"
BODY_PT = 11.5

MOVEMENTS = [
    ("Movement One: Conversation", range(1, 9)),
    ("Movement Two: Architecture", range(9, 17)),
    ("Movement Three: Emergence", range(17, 25)),
]

# ---------- helpers ----------------------------------------------------------

def set_cell(): pass

def add_field(paragraph, instr):
    """Insert a Word field (e.g. PAGE) into a paragraph run."""
    run = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fe)
    return run

def style_run(run):
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_PT)
    # ensure the east-asian/complex slots also use Garamond
    rpr = run._element.get_or_add_rPr(); rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), BODY_FONT)

def add_inline(paragraph, text, base_italic=False):
    """Render markdown **bold** and *italic* into runs."""
    # tokenize on ** first, then *
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for p in parts:
        if not p:
            continue
        italic, bold = base_italic, False
        if p.startswith('**') and p.endswith('**'):
            bold = True; p = p[2:-2]
        elif p.startswith('*') and p.endswith('*'):
            italic = True; p = p[1:-1]
        run = paragraph.add_run(p)
        run.italic = italic; run.bold = bold
        style_run(run)

def body_para(doc, text, first=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.widow_control = True
    pf.first_line_indent = Inches(0 if first else 0.30)
    pf.space_after = Pt(0)
    add_inline(p, text, base_italic=italic)
    return p

def centered(doc, text, size, bold=False, italic=False, caps=False,
             space_before=0, space_after=0, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.widow_control = True
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold; run.italic = italic
    run.font.name = BODY_FONT; run.font.size = Pt(size)
    if caps:
        rpr = run._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), '40'); rpr.append(sp)
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr(); rfonts = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rfonts.set(qn(a), BODY_FONT)
    rpr.append(rfonts)
    return p

def set_mirror_margins(doc):
    settings = doc.settings.element
    if settings.find(qn('w:mirrorMargins')) is None:
        settings.append(OxmlElement('w:mirrorMargins'))

def section_page_setup(section, restart=None, footer_num=False):
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)   # inside (mirror)
    section.right_margin = Inches(0.50)  # outside (mirror)
    section.different_first_page_header_footer = False
    sectPr = section._sectPr
    if restart is not None:
        pn = sectPr.find(qn('w:pgNumType'))
        if pn is None:
            pn = OxmlElement('w:pgNumType'); sectPr.append(pn)
        pn.set(qn('w:start'), str(restart))
    if footer_num:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(p, ' PAGE ')
        for r in p.runs:
            r.font.name = BODY_FONT; r.font.size = Pt(10)
    else:
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()

def parse_chapter(path):
    lines = path.read_text(encoding='utf-8').splitlines()
    heading = ''
    body = []
    for ln in lines:
        if ln.startswith('# ') and not heading:
            heading = ln[2:].strip()
        elif ln.strip() == '':
            body.append('')
        else:
            body.append(ln.rstrip())
    # collapse into paragraphs on blank lines
    paras, cur = [], []
    for ln in body:
        if ln == '':
            if cur: paras.append(' '.join(cur)); cur = []
        else:
            cur.append(ln)
    if cur: paras.append(' '.join(cur))
    return heading, paras

# ---------- build -----------------------------------------------------------

doc = Document()
# base Normal style
normal = doc.styles['Normal']
normal.font.name = BODY_FONT
normal.font.size = Pt(BODY_PT)
set_mirror_margins(doc)

# ----- FRONT MATTER (section 1, no page numbers) -----
sec0 = doc.sections[0]
section_page_setup(sec0, footer_num=False)

# Title page
for _ in range(4): doc.add_paragraph()
centered(doc, "PAIGE", 40, bold=True, caps=False, space_after=6)
centered(doc, "Volume One", 16, italic=True, space_after=2)
for _ in range(2): doc.add_paragraph()
centered(doc, "A Novel", 13)
for _ in range(10): doc.add_paragraph()
centered(doc, AUTHOR, 14, caps=True)

# Copyright page
doc.add_page_break()
for _ in range(2): doc.add_paragraph()
cp = [
    f"PAIGE: Volume One",
    "",
    f"Copyright © {YEAR} by {AUTHOR}",
    "",
    "All rights reserved.",
    "",
    "This is a work of fiction. Names, characters, places, institutions, and "
    "incidents are either products of the author’s imagination or used "
    "fictitiously. Any resemblance to actual persons, living or dead, events, "
    "or locales is entirely coincidental. References to Florida State "
    "University and the city of Tallahassee are used fictively and do not "
    "depict real persons, policies, or events.",
    "",
    "No part of this book may be reproduced in any form or by any electronic "
    "or mechanical means, including information storage and retrieval systems, "
    "without written permission from the author, except for the use of brief "
    "quotations in a book review.",
    "",
    "First Edition.",
    "",
    f"{AUTHOR}",
]
for line in cp:
    if line == "":
        doc.add_paragraph()
    else:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line); style_run(r); r.font.size = Pt(9.5)

# Table of Contents
doc.add_page_break()
centered(doc, "Contents", 18, space_after=14)
chapters = sorted(CH_DIR.glob('ch*.md'))
heads = {}
for path in chapters:
    n = int(re.match(r'ch(\d+)', path.name).group(1))
    h, _ = parse_chapter(path)
    heads[n] = h
for mv_title, rng in MOVEMENTS:
    centered(doc, mv_title, 11.5, italic=True, space_before=10, space_after=6)
    for n in rng:
        h = heads.get(n, '')
        # h like "Chapter Six — Ordinary Time"
        disp = h.replace('Chapter ', '')
        p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(disp); style_run(r); r.font.size = Pt(11)

# ----- BODY (section 2, restart page numbers, footer numbering) -----
new = doc.add_section(WD_SECTION.NEW_PAGE)
section_page_setup(new, restart=1, footer_num=True)

for idx, path in enumerate(chapters):
    n = int(re.match(r'ch(\d+)', path.name).group(1))
    heading, paras = parse_chapter(path)
    if ' — ' in heading:
        label, title = heading.split(' — ', 1)
    else:
        label, title = heading, ''
    # chapter opener on a new page
    opener = centered(doc, label, 12, caps=True, space_before=48, space_after=4)
    if idx != 0:
        opener.paragraph_format.page_break_before = True
    if title:
        centered(doc, title, 17, italic=False, space_after=24)
    for i, para in enumerate(paras):
        body_para(doc, para, first=(i == 0))

# ----- BACK MATTER: acknowledgements placeholder -----
ack = centered(doc, "Acknowledgements", 17, space_before=48, space_after=18)
ack.paragraph_format.page_break_before = True
body_para(doc, "[Acknowledgements to be added.]", first=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
wc = sum(len(' '.join(parse_chapter(p)[1]).split()) for p in chapters)
print(f"Wrote {OUT}  ({len(chapters)} chapters, ~{wc:,} body words)")
