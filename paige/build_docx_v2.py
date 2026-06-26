#!/usr/bin/env python3
"""Build the KDP-ready DOCX for PAIGE: Volume One — Edition 1 Version 2.

Spec: trim 5.5x8.5 | Garamond 11pt body | first-line indent 0.30in | justified
inside 0.75 / outside 0.50 / top-bottom 0.75 | gutter 0.125 | mirror margins
widow/orphan | page break before chapters | even/odd running headers
(book title on even, chapter title on odd) | centered page-number footer.
"""
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
CH_DIR = ROOT / "chapters"
OUT = ROOT / "Paige_Volume_One_Ed1_V2_2026-06-26.docx"

TITLE = "Paige: Volume One"
AUTHOR = "KS Pierre"
PUBLISHER = "PCM GenCon Group"
CONTRIBUTOR = "Claude"
VERSION = "Edition 1 Version 2"
BODY_FONT = "Garamond"
BODY_PT = 11.0

MOVEMENTS = [
    ("Movement One: Conversation", range(1, 9)),
    ("Movement Two: Architecture", range(9, 17)),
    ("Movement Three: Emergence", range(17, 25)),
]

def add_field(paragraph, instr):
    run = paragraph.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    run._r.append(fb); run._r.append(it); run._r.append(fe)
    return run

def garamond(run, size=None, bold=None, italic=None, caps=False, color=None):
    run.font.name = BODY_FONT
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'): rfonts.set(qn(a), BODY_FONT)
    if caps:
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), '40'); rpr.append(sp)

def inline(paragraph, text, base_italic=False):
    for p in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if not p: continue
        italic, bold = base_italic, False
        if p.startswith('**') and p.endswith('**'): bold = True; p = p[2:-2]
        elif p.startswith('*') and p.endswith('*'): italic = True; p = p[1:-1]
        run = paragraph.add_run(p); garamond(run, italic=italic, bold=bold)

def body_para(doc, text, first=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    pf.widow_control = True
    pf.first_line_indent = Inches(0 if first else 0.30)
    pf.space_after = Pt(0); pf.space_before = Pt(0)
    inline(p, text)
    return p

def centered(doc, text, size, bold=False, italic=False, caps=False,
             sb=0, sa=0, heading=False):
    p = doc.add_paragraph()
    if heading:
        p.style = doc.styles['Heading 1']
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.widow_control = True; pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.keep_with_next = True
    run = p.add_run(text.upper() if caps else text)
    garamond(run, size=size, bold=bold, italic=italic, caps=caps,
             color=RGBColor(0, 0, 0))
    return p

def page_setup(section, gutter=True):
    section.page_width = Inches(5.5); section.page_height = Inches(8.5)
    section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75); section.right_margin = Inches(0.50)
    if gutter: section.gutter = Inches(0.125)

def footer_pagenum(footer):
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' PAGE ')
    for r in p.runs: garamond(r, size=10)

def header_text(header, text):
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); garamond(run, size=9.5, italic=True)

def header_styleref(header):
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' STYLEREF "Heading 1" \\* MERGEFORMAT ')
    for r in p.runs: garamond(r, size=9.5, italic=True)

def empty_hf(part):
    part.is_linked_to_previous = False
    if not part.paragraphs: part.add_paragraph()

def parse_chapter(path):
    heading = ''; body = []
    for ln in path.read_text(encoding='utf-8').splitlines():
        if ln.startswith('# ') and not heading: heading = ln[2:].strip()
        elif ln.strip() == '': body.append('')
        else: body.append(ln.rstrip())
    paras, cur = [], []
    for ln in body:
        if ln == '':
            if cur: paras.append(' '.join(cur)); cur = []
        else: cur.append(ln)
    if cur: paras.append(' '.join(cur))
    return heading, paras

# ---------------- build ----------------
doc = Document()
normal = doc.styles['Normal']; normal.font.name = BODY_FONT; normal.font.size = Pt(BODY_PT)

# settings: mirror margins + even/odd headers
settings = doc.settings.element
for tag in ('w:mirrorMargins', 'w:evenAndOddHeaders'):
    if settings.find(qn(tag)) is None: settings.append(OxmlElement(tag))

# FRONT MATTER section (no headers/footers, no page numbers)
sec0 = doc.sections[0]; page_setup(sec0, gutter=True)
for hf in (sec0.header, sec0.even_page_header, sec0.first_page_header,
           sec0.footer, sec0.even_page_footer, sec0.first_page_footer):
    empty_hf(hf)
sec0.different_first_page_header_footer = False

for _ in range(4): doc.add_paragraph()
centered(doc, "PAIGE", 40, bold=True, sa=6)
centered(doc, "Volume One", 16, italic=True, sa=2)
for _ in range(2): doc.add_paragraph()
centered(doc, "A Novel", 13)
for _ in range(10): doc.add_paragraph()
centered(doc, AUTHOR, 14, caps=True)

# Copyright page
doc.add_page_break()
for _ in range(2): doc.add_paragraph()
cp = [
    "Paige: Volume One", "",
    f"Copyright © 2026 by {AUTHOR}", "",
    "All rights reserved.", "",
    "This is a work of fiction. Names, characters, places, institutions, and "
    "incidents are either products of the author’s imagination or used "
    "fictitiously. Any resemblance to actual persons, living or dead, events, or "
    "locales is entirely coincidental. References to Florida State University and "
    "the city of Tallahassee are used fictively and do not depict real persons, "
    "policies, or events.", "",
    "No part of this book may be reproduced in any form or by any electronic or "
    "mechanical means, including information storage and retrieval systems, "
    "without written permission from the author, except for the use of brief "
    "quotations in a book review.", "",
    f"Published by {PUBLISHER}.", "",
    "First Edition.", "",
    AUTHOR,
]
for line in cp:
    if line == "": doc.add_paragraph()
    else:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        garamond(p.add_run(line), size=9.5)

# Contents
doc.add_page_break()
centered(doc, "Contents", 18, bold=True, sa=14)
chapters = sorted(CH_DIR.glob('ch*.md'))
heads = {}
for path in chapters:
    n = int(re.match(r'ch(\d+)', path.name).group(1))
    heads[n] = parse_chapter(path)[0]
if 0 in heads:
    p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6); garamond(p.add_run(heads[0]), size=11)
for mv_title, rng in MOVEMENTS:
    centered(doc, mv_title, 11.5, italic=True, sb=10, sa=6)
    for n in rng:
        h = heads.get(n, '')
        if ' — ' in h:
            lab, ti = h.split(' — ', 1); disp = f"{lab.replace('Chapter ', '')}: {ti}"
        else:
            disp = h.replace('Chapter ', '')
        p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0); garamond(p.add_run(disp), size=11)

# BODY section: restart page numbers at 1, even/odd headers, footer page number
body = doc.add_section(WD_SECTION.NEW_PAGE); page_setup(body, gutter=True)
sectPr = body._sectPr
pn = OxmlElement('w:pgNumType'); pn.set(qn('w:start'), '1'); sectPr.append(pn)
header_styleref(body.header)               # odd pages: chapter title
header_text(body.even_page_header, TITLE)  # even pages: book title
empty_hf(body.first_page_header)
footer_pagenum(body.footer); footer_pagenum(body.even_page_footer)
empty_hf(body.first_page_footer)

for idx, path in enumerate(chapters):
    heading, paras = parse_chapter(path)
    label, title = (heading.split(' — ', 1) + [''])[:2] if ' — ' in heading else (heading, '')
    if title:
        op = centered(doc, label, 12, caps=True, sb=48, sa=4)
        centered(doc, title, 17, bold=True, sa=24, heading=True)  # Heading 1 -> running header
    else:
        op = centered(doc, label, 18, bold=True, caps=True, sb=48, sa=24, heading=True)
    if idx != 0: op.paragraph_format.page_break_before = True
    for i, para in enumerate(paras):
        body_para(doc, para, first=(i == 0))

# Acknowledgements
ack = centered(doc, "Acknowledgements", 17, bold=True, sb=48, sa=18, heading=True)
ack.paragraph_format.page_break_before = True
body_para(doc, "[Acknowledgements to be added.]", first=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- metadata ----
cp = doc.core_properties
cp.title = TITLE
cp.author = AUTHOR
cp.last_modified_by = AUTHOR
cp.category = "Literary Fiction"
cp.comments = VERSION
cp.version = VERSION
cp.created = datetime(2026, 6, 26, tzinfo=timezone.utc)
cp.modified = datetime(2026, 6, 26, tzinfo=timezone.utc)
cp.keywords = "novel, literary fiction"

doc.save(OUT)

# ---- custom properties (Publisher, Contributor, Version, Date) + app Company ----
import zipfile, shutil, tempfile, os
custom_xml = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
    'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="Publisher">'
    f'<vt:lpwstr>{PUBLISHER}</vt:lpwstr></property>'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="3" name="Contributor">'
    f'<vt:lpwstr>{CONTRIBUTOR}</vt:lpwstr></property>'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="4" name="Creator">'
    f'<vt:lpwstr>{AUTHOR}</vt:lpwstr></property>'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="5" name="Version">'
    f'<vt:lpwstr>{VERSION}</vt:lpwstr></property>'
    '<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="6" name="Date">'
    '<vt:lpwstr>26 June 2026</vt:lpwstr></property>'
    '</Properties>'
)
ct_add = '<Override PartName="/docProps/custom.xml" ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/>'
rel_add = ('<Relationship Id="rIdCustom1" '
           'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties" '
           'Target="docProps/custom.xml"/>')

tmp = OUT.with_suffix('.tmp.docx')
with zipfile.ZipFile(OUT, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
    for item in zin.namelist():
        data = zin.read(item)
        if item == '[Content_Types].xml':
            data = data.replace(b'</Types>', (ct_add + '</Types>').encode())
        elif item == '_rels/.rels':
            data = data.replace(b'</Relationships>', (rel_add + '</Relationships>').encode())
        elif item == 'docProps/app.xml':
            txt = data.decode('utf-8')
            if '<Company>' not in txt:
                txt = txt.replace('</Properties>', f'<Company>{PUBLISHER}</Company></Properties>')
            data = txt.encode('utf-8')
        zout.writestr(item, data)
    zout.writestr('docProps/custom.xml', custom_xml)
shutil.move(tmp, OUT)

wc = sum(len(' '.join(parse_chapter(p)[1]).split()) for p in chapters)
print(f"Wrote {OUT.name}  ({len(chapters)} sections incl. prologue, ~{wc:,} body words)")
